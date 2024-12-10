import win32com.client
import os
from datetime import datetime, time
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkcalendar import DateEntry
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from collections import defaultdict

class FinNewsApp:
    """
    Financial News Processor
    
    Requirements:
    - pywin32 (for Outlook access): pip install pywin32
    - python-docx (for Word export): pip install python-docx
    - tkcalendar (for date selection): pip install tkcalendar
    
    Usage:
    1. Click "连接Outlook" to connect to Outlook.
    2. Select a mailbox folder from the tree.
    3. Enter optional keywords and select date range.
    4. Click "搜索" to retrieve and classify emails.
    5. Check preview and classification results.
    6. Choose an export path and click "保存到Word" to export.
    """
    def __init__(self, master):
        self.master = master
        self.master.title("Financial News Processor")
        self.master.geometry("900x700")
        self.outlook = None

        # Define categories and their keywords
        self.classification_keywords = {
            "中国": ["中国", "中国央行", "财政部", "蓝佛安", "中央"],
            "美国": ["美国", "美联储", "拜登", "特朗普"],
        }

        self.selected_folder = None
        self.processed_messages = []
        self.news_dict = {}  # Will store categorized news

        self.setup_ui()

    def setup_ui(self):
        style = ttk.Style()
        style.theme_use('clam')

        # Main container
        main_frame = ttk.Frame(self.master, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Connection status
        status_frame = ttk.LabelFrame(main_frame, text="连接状态", padding="5")
        status_frame.pack(fill=tk.X, pady=(0, 10))

        self.status_var = tk.StringVar(value="未连接")
        ttk.Label(status_frame, textvariable=self.status_var).pack(side=tk.LEFT)
        ttk.Button(status_frame, text="连接Outlook", command=self.connect_outlook).pack(side=tk.LEFT, padx=5)
        self.account_label = ttk.Label(status_frame, text="")
        self.account_label.pack(side=tk.LEFT, padx=10)

        # Folder selection frame
        folder_frame = ttk.LabelFrame(main_frame, text="文件夹选择", padding="5")
        folder_frame.pack(fill=tk.X, pady=(0, 10))

        tree_frame = ttk.Frame(folder_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        self.folder_tree = ttk.Treeview(tree_frame, height=6)
        self.folder_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.folder_tree.heading("#0", text="邮件文件夹", anchor="w")

        folder_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.folder_tree.yview)
        folder_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.folder_tree.configure(yscrollcommand=folder_scrollbar.set)

        # Search options
        search_frame = ttk.LabelFrame(main_frame, text="搜索选项", padding="5")
        search_frame.pack(fill=tk.X, pady=(0, 10))

        # Keywords
        ttk.Label(search_frame, text="关键词(可选,逗号分隔):").pack(side=tk.LEFT, padx=(0, 5))
        self.keyword_entry = ttk.Entry(search_frame, width=40)
        self.keyword_entry.pack(side=tk.LEFT, padx=5)

        # Date range
        date_frame = ttk.Frame(search_frame)
        date_frame.pack(side=tk.LEFT, padx=5)

        ttk.Label(date_frame, text="开始日期:").pack(side=tk.LEFT)
        self.start_date = DateEntry(date_frame, width=12)
        self.start_date.pack(side=tk.LEFT, padx=5)

        ttk.Label(date_frame, text="结束日期:").pack(side=tk.LEFT)
        self.end_date = DateEntry(date_frame, width=12)
        self.end_date.pack(side=tk.LEFT, padx=5)

        # Export options
        export_frame = ttk.LabelFrame(main_frame, text="导出选项", padding="5")
        export_frame.pack(fill=tk.X, pady=(0, 10))

        self.export_path_label = ttk.Label(export_frame, text="导出路径:")
        self.export_path_label.pack(side=tk.LEFT)
        
        self.export_path_entry = ttk.Entry(export_frame, width=50)
        self.export_path_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(export_frame, text="选择路径", command=self.browse_export_path).pack(side=tk.LEFT, padx=5)

        # Buttons section
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(button_frame, text="搜索", command=self.search).pack(side=tk.LEFT, padx=5)
        self.save_button = ttk.Button(button_frame, text="保存到Word", command=self.save_to_word, state=tk.DISABLED)
        self.save_button.pack(side=tk.LEFT, padx=5)
        
        # Clear results button
        ttk.Button(button_frame, text="清除结果", command=self.clear_results).pack(side=tk.LEFT, padx=5)

        # Results display
        result_frame = ttk.LabelFrame(main_frame, text="预览", padding="5")
        result_frame.pack(fill=tk.BOTH, expand=True)

        self.result_text = tk.Text(result_frame, wrap=tk.WORD)
        self.result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(result_frame, orient=tk.VERTICAL, command=self.result_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.result_text.configure(yscrollcommand=scrollbar.set)

        # Status bar at bottom
        self.bottom_status_var = tk.StringVar(value="准备就绪")
        bottom_status = ttk.Label(self.master, textvariable=self.bottom_status_var, relief=tk.SUNKEN, anchor='w')
        bottom_status.pack(side=tk.BOTTOM, fill=tk.X)

    def connect_outlook(self):
        try:
            self.outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")
            self.status_var.set("已连接")
            account = self.outlook.CurrentUser.Name
            self.account_label.config(text=f"账户: {account}")
            self.show_folders()
            self.bottom_status_var.set("Outlook连接成功")
        except Exception as e:
            messagebox.showerror("错误", f"连接Outlook失败: {str(e)}")
            self.bottom_status_var.set("连接Outlook失败")

    def show_folders(self):
        self.folder_tree.delete(*self.folder_tree.get_children())
        try:
            inbox = self.outlook.GetDefaultFolder(6)  # 6 = Inbox
            self.add_folder_to_tree("", inbox)
            self.bottom_status_var.set("文件夹已加载")
        except Exception as e:
            self.bottom_status_var.set(f"加载文件夹时出错: {e}")

    def add_folder_to_tree(self, parent, folder):
        try:
            folder_id = self.folder_tree.insert(parent, "end", text=folder.Name, values=(folder.EntryID, folder.StoreID))
            for subfolder in folder.Folders:
                self.add_folder_to_tree(folder_id, subfolder)
        except Exception:
            # Silently ignore certain problematic folders
            pass

    def get_selected_folder(self):
        selected_id = self.folder_tree.selection()
        if not selected_id:
            return None
        folder_entry_id, store_id = self.folder_tree.item(selected_id, 'values')
        folder = self.outlook.GetFolderFromID(folder_entry_id, store_id)
        return folder

    def clean_text(self, text):
        if not text:
            return ""
        
        lines = text.splitlines()
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Only keep lines starting with * and not containing '今晨央行'
            if line.startswith('*') and '今晨央行' not in line:
                cleaned_lines.append(line)
        
        content = '\n'.join(cleaned_lines)
        # Remove multiple blank lines
        content = re.sub(r'\n\s*\n', '\n', content)
        
        return content.strip()

    def search(self):
        if not self.outlook:
            messagebox.showerror("错误", "请先点击'连接Outlook'按钮进行连接。")
            return

        selected_folder = self.get_selected_folder()
        if not selected_folder:
            messagebox.showerror("错误", "请选择要搜索的文件夹后再试。")
            return

        self.bottom_status_var.set("正在搜索，请稍候...")
        self.master.update_idletasks()

        try:
            messages = selected_folder.Items
            messages.Sort("[ReceivedTime]", True)

            # Get keywords and dates
            keywords = [k.strip() for k in self.keyword_entry.get().split(',') if k.strip()]
            start_date = self.start_date.get_date()
            end_date = self.end_date.get_date()

            start_datetime = datetime.combine(start_date, time.min)
            end_datetime = datetime.combine(end_date, time.max)
            start_date_str = start_datetime.strftime('%m/%d/%Y')
            end_date_str = end_datetime.strftime('%m/%d/%Y')

            # Date filter
            date_filter = f"[ReceivedTime] >= '{start_date_str}' AND [ReceivedTime] <= '{end_date_str}'"
            messages = messages.Restrict(date_filter)

            # Keyword filter if needed
            if keywords:
                keyword_conditions = []
                for keyword in keywords:
                    keyword_conditions.append(f'@SQL="urn:schemas:httpmail:subject" LIKE \'%{keyword}%\'')
                keyword_filter = " OR ".join(keyword_conditions)
                messages = messages.Restrict(keyword_filter)

            self.processed_messages.clear()
            for msg in messages:
                if msg.Body:
                    cleaned_body = self.clean_text(msg.Body)
                    if cleaned_body:
                        self.processed_messages.append({
                            'body': cleaned_body,
                            'time': msg.ReceivedTime,
                            'subject': msg.Subject
                        })

            self.processed_messages.sort(key=lambda x: x['time'])

            self.display_preview(self.processed_messages)

            # Now classify and extract by building doc_content-like structure
            doc_content = ""
            for msg in self.processed_messages:
                doc_content += f"时间: {msg['time'].strftime('%Y-%m-%d %H:%M:%S')}\n"
                doc_content += msg['body'] + "\n\n"

            self.news_dict = self.extract_news(doc_content)
            self.display_news(self.news_dict)

            # If we have processed messages, allow saving even if classification failed
            self.save_button.config(state=tk.NORMAL if self.processed_messages else tk.DISABLED)

            if not self.processed_messages:
                messagebox.showinfo("结果", "未找到符合条件的邮件。请调整关键词或日期范围重试。")
                self.bottom_status_var.set("未找到邮件")
            else:
                self.bottom_status_var.set("搜索完成")

        except Exception as e:
            messagebox.showerror("错误", f"搜索过程中出错: {str(e)}")
            self.bottom_status_var.set("搜索出错")

    def save_to_word(self):
        if not self.processed_messages:
            messagebox.showerror("错误", "没有可保存的搜索结果，请先搜索邮件。")
            return

        if not self.export_path_entry.get():
            messagebox.showerror("错误", "请选择有效的导出路径后再保存。")
            return

        if not os.path.isdir(self.export_path_entry.get()):
            messagebox.showerror("错误", "导出路径不存在，请选择有效的路径。")
            return

        self.bottom_status_var.set("正在保存，请稍候...")
        self.master.update_idletasks()

        try:
            self.export_to_word(self.news_dict)
            self.bottom_status_var.set("保存完成")
        except Exception as e:
            messagebox.showerror("错误", f"保存文档时出错: {str(e)}")
            self.bottom_status_var.set("保存出错")

    def browse_export_path(self):
        path = filedialog.askdirectory()
        if path:
            self.export_path_entry.delete(0, tk.END)
            self.export_path_entry.insert(0, path)

    def export_to_word(self, news_dict):
        doc = Document()
        
        # Add title
        title = doc.add_heading('邮件内容汇总', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set default font and style
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)

        # Add metadata
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"搜索时间范围: {self.start_date.get_date().strftime('%Y-%m-%d')} 至 {self.end_date.get_date().strftime('%Y-%m-%d')}")
        doc.add_paragraph("=" * 50)

        # Add content by category
        if news_dict:
            for category, news_list in news_dict.items():
                doc.add_heading(f"{category} News", level=1)
                # Sort by date
                sorted_news = sorted(news_list, key=lambda x: x['date'] if x['date'] else datetime.min)
                for entry in sorted_news:
                    date_str = entry['date'].strftime('%Y-%m-%d') if entry['date'] else "未知日期"
                    p = doc.add_paragraph()
                    p.add_run(f"时间: {date_str}\n").bold = True
                    p.add_run(f"{entry['news']}\n")

                doc.add_paragraph("_" * 30)
        else:
            # If no classified news, just list original messages
            doc.add_heading("未分类信息", level=1)
            for msg in self.processed_messages:
                p = doc.add_paragraph()
                p.add_run(f"时间: {msg['time'].strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
                p.add_run(f"{msg['body']}\n")
            doc.add_paragraph("_" * 30)

        # Save document
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        doc_path = os.path.join(self.export_path_entry.get(), f'邮件汇总_{timestamp}.docx')
        doc.save(doc_path)
        messagebox.showinfo("完成", f"文档已保存至:\n{doc_path}")

    def display_preview(self, messages):
        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, f"共找到 {len(messages)} 封邮件(原始内容)\n\n")
        
        for i, msg in enumerate(messages, 1):
            self.result_text.insert(tk.END, f"{i}. 主题: {msg['subject']}\n")
            self.result_text.insert(tk.END, f"时间: {msg['time'].strftime('%Y-%m-%d %H:%M:%S')}\n")
            preview = msg['body'].replace('\n', ' ').strip()
            if len(preview) > 500:
                preview = preview[:500] + "..."
            self.result_text.insert(tk.END, f"内容预览: {preview}\n")
            self.result_text.insert(tk.END, "-" * 50 + "\n\n")

    def classify_news(self, news):
        """Classify the news into China, US, or International based on keywords."""
        for category, kw_list in self.classification_keywords.items():
            if any(keyword in news for keyword in kw_list):
                return category
        return "国际"

    def extract_news(self, doc_content):
        # Capture date and news entries
        date_pattern = re.compile(r"时间:\s*(\d{4}-\d{2}-\d{2})\s*\d{2}:\d{2}:\d{2}")
        # Updated pattern: no trailing \n, line-start anchor
        news_pattern = re.compile(r"^\*\s*(.*)")

        news_by_category = defaultdict(list)
        current_date = None

        lines = doc_content.split("\n")
        for line in lines:
            # Match date
            date_match = date_pattern.search(line)
            if date_match:
                current_date = datetime.strptime(date_match.group(1), "%Y-%m-%d")

            # Match news lines
            news_matches = news_pattern.findall(line)
            for news in news_matches:
                category = self.classify_news(news)
                news_by_category[category].append({"date": current_date, "news": news})
            
        return news_by_category

    def display_news(self, news_dict):
        self.result_text.insert(tk.END, "\n\n分类后新闻:\n")
        if not news_dict:
            self.result_text.insert(tk.END, "无分类新闻可展示。\n")
            return
        for category, news_list in news_dict.items():
            self.result_text.insert(tk.END, f"\n--- {category} News ---\n")
            sorted_news = sorted(news_list, key=lambda x: x['date'] if x['date'] else datetime.min)
            for entry in sorted_news:
                date_str = entry['date'].strftime('%Y-%m-%d') if entry['date'] else "未知日期"
                self.result_text.insert(tk.END, f"{date_str}: {entry['news']}\n")

    def clear_results(self):
        self.processed_messages.clear()
        self.news_dict.clear()
        self.result_text.delete(1.0, tk.END)
        self.save_button.config(state=tk.DISABLED)
        self.bottom_status_var.set("结果已清除")

def main():
    root = tk.Tk()
    app = FinNewsApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
